import docx
import logging
import argparse
import re
import sys

logging.basicConfig(level=logging.DEBUG)

# Set up logging
def get_arg():
    """ Takes nothing
Purpose: Gets arguments from command line
Returns: Argument's values
"""
    parser = argparse.ArgumentParser()
    # CLI Version
    # parser.add_argument("-d","--debug",dest="debug",action="store_true",help="Turn on debugging",default=False)
    parser.add_argument("-d","--debug",dest="debug",action="store_false",help="Turn on debugging",default=True)
    # File version
    parser.add_argument("-f","--file",dest="file", help="Name of the Word Doc.")
    parser.add_argument("-o","--output",dest="output", help="Name of the file to output the results.")
    options = parser.parse_args()

    if not options.output:
        options.output = "output.csv"
    if not options.file:
        logging.error("Please provide a file name.")
        sys.exit()

    if options.debug:
        logging.basicConfig(level=logging.DEBUG)
        global DEBUG
        DEBUG = True
    else:
        logging.basicConfig(level=logging.INFO)
    return options


def tableSearch(doc):
    """ Takes doc file
Purpose: Scrapes all tables and pulls back text
Returns: Dictionary of tables
"""
    tables = {}
    table = 0
    x = 0
    y = 0

    while table != len(doc.tables):
        logging.debug("Table {} Row: {}".format(table, len(doc.tables[table].rows)))
        logging.debug("Table {} Column: {}".format(table, len(doc.tables[table].columns)))
        logging.debug("Table {}".format(table))
        table_test = doc.tables[table]
        run = doc.add_paragraph().add_run()
        font = run.font
        while x != len(table_test.rows):
            while y != len(table_test.columns):
                logging.debug("Table: {}, X: {}, Y: {}\n{}\n".format(table, x,y,table_test.cell(x,y).text))
                tables[str(table)+"."+str(x)+"."+str(y)] = {"row": x, "column": y, "text": table_test.cell(x,y).text}
                y += 1
            x += 1
            y = 0
        x = 0
        y = 0
        table += 1

    return tables


# Reading the word document
def wordDocx(file):
    """ Takes file
Purpose: Reads the Word Doc
Returns: Nothing
"""
    logging.debug("File: {}".format(file))
    doc = docx.Document(file)
    fullText = []

    for para in doc.paragraphs:
        logging.debug("Paragraph: {}".format(para.text))
        fullText.append(para.text)

    tableInfo = tableSearch(doc)
    return [fullText, tableInfo]


def main():
    options = get_arg()
    logging.debug("Options: {}".format(options))
    info = wordDocx(options.file)
    logging.debug("Info: {}".format(info))

if __name__ == '__main__':
    main()
